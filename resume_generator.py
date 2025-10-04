#!/usr/bin/env python3

"""
Module 3: Command-Line Resume Generator
Takes 2 JSON files (resume data + job data) and generates tailored PDF/Word documents

Usage:
python resume_generator.py resume.json job.json --output resume_tailored --format pdf
python resume_generator.py resume_data.json job_requirements.json --format docx --api-key YOUR_KEY
"""

import json
import argparse
import sys
import os
from pathlib import Path
from typing import Dict, Any, List
import logging
from datetime import datetime

# Import our core components
try:
    import requests
except ImportError:
    print("❌ ERROR: requests not installed. Run: pip install requests")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class CommandLineResumeGenerator:
    """Command-line resume generator using Groq API"""
    
    def __init__(self, api_key: str = None):
        # API key can be provided via command line or hardcoded here
        self.api_key = api_key or "YOUR_GROQ_API_KEY_HERE"  # Replace with your key
        self.model = "qwen-3-32b"
        self.base_url = "https://api.groq.com/openai/v1/chat/completions"
        
        if self.api_key == "YOUR_GROQ_API_KEY_HERE":
            print("⚠️ WARNING: Using default API key placeholder. Set your key with --api-key or edit the code.")
    
    def load_json_file(self, file_path: str) -> Dict[str, Any]:
        """Load and validate JSON file"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            logger.info(f"✅ Loaded JSON file: {file_path}")
            return data
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON in {file_path}: {str(e)}")
        except Exception as e:
            raise Exception(f"Error reading {file_path}: {str(e)}")
    
    def create_tailoring_prompt(self, resume_data: Dict[str, Any], job_data: Dict[str, Any]) -> str:
        """Create optimized prompt for resume tailoring"""
        # Extract key information
        candidate_name = resume_data.get("personal_info", {}).get("name", "Candidate")
        job_title = job_data.get("job_title", "Target Position")
        
        # Get skills from job
        required_skills = job_data.get("required_skills", [])
        preferred_skills = job_data.get("preferred_skills", [])
        responsibilities = job_data.get("key_responsibilities", [])
        
        return f"""You are an expert resume writer. Create a tailored, ATS-friendly resume that aligns the candidate's experience with the job requirements.

CANDIDATE DATA:
{json.dumps(resume_data, indent=2)}

JOB REQUIREMENTS:
Position: {job_title}
Required Skills: {', '.join(required_skills[:15])}
Preferred Skills: {', '.join(preferred_skills[:10])}
Key Responsibilities: {'; '.join(responsibilities[:5])}

INSTRUCTIONS:
1. Reorganize experience to highlight job-relevant achievements
2. Integrate target keywords naturally throughout the resume
3. Emphasize quantified achievements that align with job responsibilities
4. Ensure ATS-friendly formatting with clear sections
5. Write compelling content while maintaining authenticity

Return ONLY this JSON structure:
{{
  "personal_info": {{
    "name": "Full Name",
    "email": "email@domain.com",
    "phone": "phone number",
    "location": "City, State",
    "linkedin": "LinkedIn URL if available",
    "portfolio": "Portfolio URL if available"
  }},
  "professional_summary": "Compelling 3-4 sentence summary tailored to the target job",
  "core_competencies": [
    "Skill 1 (prioritized for job relevance)",
    "Skill 2",
    "Up to 12 most relevant skills"
  ],
  "professional_experience": [
    {{
      "position": "Job Title",
      "company": "Company Name",
      "location": "City, State",
      "duration": "Start - End",
      "achievements": [
        "• Tailored achievement showing impact relevant to target job",
        "• Quantified result with metrics (numbers, percentages, amounts)",
        "• Achievement demonstrating skills needed for target position"
      ]
    }}
  ],
  "education": [
    {{
      "degree": "Degree Type",
      "field": "Field of Study",
      "institution": "School Name",
      "graduation_year": "Year"
    }}
  ],
  "technical_skills": {{
    "programming_languages": ["Languages relevant to job"],
    "frameworks_tools": ["Frameworks/tools matching job"],
    "databases": ["Database technologies if applicable"],
    "other_technical": ["Other relevant technical skills"]
  }},
  "projects": [
    {{
      "name": "Project Name",
      "description": "Description emphasizing job-relevant skills",
      "technologies": ["Technologies matching job requirements"]
    }}
  ],
  "certifications": [
    {{
      "name": "Certification Name",
      "issuer": "Issuing Organization",
      "date": "Date if recent/relevant"
    }}
  ]
}}

Return ONLY valid JSON. Focus on job relevance and ATS optimization."""
    
    def tailor_resume(self, resume_data: Dict[str, Any], job_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate tailored resume using Groq API"""
        if self.api_key == "YOUR_GROQ_API_KEY_HERE":
            return {
                "success": False,
                "error": "API key not set. Use --api-key option or edit the code."
            }
        
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            prompt = self.create_tailoring_prompt(resume_data, job_data)
            payload = {
                "model": self.model,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are an expert resume writer specializing in ATS optimization. Return only valid JSON."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "temperature": 0.3,
                "max_tokens": 4000,
                "stream": False
            }
            
            print(f"🤖 Tailoring resume with Groq {self.model}...")
            response = requests.post(self.base_url, headers=headers, json=payload, timeout=60)
            
            if response.status_code == 200:
                response_data = response.json()
                content = response_data["choices"][0]["message"]["content"]
                
                # Extract JSON
                json_data = self._extract_json(content)
                if json_data:
                    try:
                        tailored_resume = json.loads(json_data)
                        print("✅ Resume tailoring completed successfully")
                        return {
                            "success": True,
                            "tailored_resume": tailored_resume,
                            "model": self.model
                        }
                    except json.JSONDecodeError as e:
                        return {
                            "success": False,
                            "error": f"Invalid JSON response: {str(e)}"
                        }
                else:
                    return {
                        "success": False,
                        "error": "No valid JSON found in response"
                    }
            else:
                return {
                    "success": False,
                    "error": f"API error: {response.status_code} - {response.text}"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": f"Resume tailoring failed: {str(e)}"
            }
    
    def _extract_json(self, content: str) -> str:
        """Extract JSON from response content"""
        # Look for JSON in code blocks
        if "```json" in content:
            start = content.find("```json") + 7
            end = content.find("```", start)
            if end != -1:
                return content[start:end].strip()
        
        # Find JSON between braces
        start = content.find('{')
        end = content.rfind('}') + 1
        if start != -1 and end > start:
            return content[start:end]
        
        # Try entire content
        content = content.strip()
        if content.startswith('{') and content.endswith('}'):
            return content
        
        return ""
    
    def generate_docx(self, resume_data: Dict[str, Any], output_path: str) -> Dict[str, Any]:
        """Generate DOCX document"""
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx not available. Install with: pip install python-docx"
            }
        
        try:
            doc = Document()
            
            # Set document styles
            self._setup_docx_styles(doc)
            
            # Add resume sections
            self._add_personal_section(doc, resume_data.get("personal_info", {}))
            self._add_summary_section(doc, resume_data.get("professional_summary", ""))
            self._add_skills_section(doc, resume_data.get("core_competencies", []))
            self._add_experience_section(doc, resume_data.get("professional_experience", []))
            self._add_education_section(doc, resume_data.get("education", []))
            self._add_technical_skills_section(doc, resume_data.get("technical_skills", {}))
            self._add_projects_section(doc, resume_data.get("projects", []))
            self._add_certifications_section(doc, resume_data.get("certifications", []))
            
            # Save document
            doc.save(output_path)
            file_size = os.path.getsize(output_path) / 1024
            
            print(f"✅ DOCX generated: {output_path} ({file_size:.1f} KB)")
            
            return {
                "success": True,
                "output_path": output_path,
                "format": "DOCX",
                "file_size_kb": file_size
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX generation failed: {str(e)}"
            }
    
    def generate_pdf(self, resume_data: Dict[str, Any], output_path: str) -> Dict[str, Any]:
        """Generate PDF using Pandoc"""
        try:
            # Generate markdown content
            markdown_content = self._generate_markdown(resume_data)
            
            # Create temporary markdown file
            temp_md = output_path.replace('.pdf', '_temp.md')
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            # Convert to PDF using Pandoc
            import subprocess
            pandoc_cmd = [
                'pandoc', temp_md, '-o', output_path,
                '--pdf-engine=xelatex',
                '-V', 'geometry:margin=1in',
                '-V', 'fontsize=11pt',
                '-V', 'linestretch=1.1'
            ]
            
            result = subprocess.run(pandoc_cmd, capture_output=True, text=True, timeout=30)
            
            # Clean up temp file
            if os.path.exists(temp_md):
                os.remove(temp_md)
            
            if result.returncode == 0 and os.path.exists(output_path):
                file_size = os.path.getsize(output_path) / 1024
                print(f"✅ PDF generated: {output_path} ({file_size:.1f} KB)")
                
                return {
                    "success": True,
                    "output_path": output_path,
                    "format": "PDF",
                    "file_size_kb": file_size
                }
            else:
                return {
                    "success": False,
                    "error": f"Pandoc conversion failed: {result.stderr}"
                }
                
        except FileNotFoundError:
            return {
                "success": False,
                "error": "Pandoc not found. Install from: https://pandoc.org/installing.html"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF generation failed: {str(e)}"
            }
    
    def _setup_docx_styles(self, doc):
        """Setup document styles"""
        # Normal style
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
    
    def _add_personal_section(self, doc, personal_info):
        """Add personal information"""
        if personal_info.get("name"):
            name_para = doc.add_paragraph(personal_info["name"])
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_para.runs[0].font.size = Pt(16)
            name_para.runs[0].font.bold = True
        
        # Contact info
        contact_parts = []
        for field in ['email', 'phone', 'location']:
            if personal_info.get(field):
                contact_parts.append(personal_info[field])
        
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
    
    def _add_summary_section(self, doc, summary):
        """Add professional summary"""
        if summary:
            heading = doc.add_paragraph("PROFESSIONAL SUMMARY")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)
            doc.add_paragraph(summary)
            doc.add_paragraph()
    
    def _add_skills_section(self, doc, competencies):
        """Add core competencies"""
        if competencies:
            heading = doc.add_paragraph("CORE COMPETENCIES")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)
            skills_text = " • ".join(competencies)
            doc.add_paragraph(skills_text)
            doc.add_paragraph()
    
    def _add_experience_section(self, doc, experience):
        """Add professional experience"""
        if experience:
            heading = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)
            
            for job in experience:
                # Job header
                job_title = f"{job.get('position', '')} | {job.get('company', '')}"
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(job_title)
                job_run.bold = True
                
                # Details
                details = f"{job.get('duration', '')} | {job.get('location', '')}"
                doc.add_paragraph(details)
                
                # Achievements
                for achievement in job.get('achievements', []):
                    doc.add_paragraph(achievement, style='List Bullet')
                
                doc.add_paragraph()  # Spacing between jobs
    
    def _add_education_section(self, doc, education):
        """Add education"""
        if education:
            heading = doc.add_paragraph("EDUCATION")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)
            
            for edu in education:
                edu_text = f"{edu.get('degree', '')} in {edu.get('field', '')}"
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(edu_text)
                edu_run.bold = True
                
                institution = f"{edu.get('institution', '')} | {edu.get('graduation_year', '')}"
                doc.add_paragraph(institution)
            
            doc.add_paragraph()
    
    def _add_technical_skills_section(self, doc, tech_skills):
        """Add technical skills"""
        if tech_skills:
            heading = doc.add_paragraph("TECHNICAL SKILLS")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)
            
            for category, skills in tech_skills.items():
                if skills:
                    category_name = category.replace('_', ' ').title()
                    skills_text = f"{category_name}: {', '.join(skills)}"
                    doc.add_paragraph(skills_text)
            
            doc.add_paragraph()
    
    def _add_projects_section(self, doc, projects):
        """Add projects"""
        if projects:
            heading = doc.add_paragraph("PROJECTS")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)
            
            for project in projects:
                project_para = doc.add_paragraph()
                project_run = project_para.add_run(project.get('name', ''))
                project_run.bold = True
                
                if project.get('description'):
                    doc.add_paragraph(project['description'])
                
                if project.get('technologies'):
                    tech_text = f"Technologies: {', '.join(project['technologies'])}"
                    doc.add_paragraph(tech_text)
                
                doc.add_paragraph()
    
    def _add_certifications_section(self, doc, certifications):
        """Add certifications"""
        if certifications:
            heading = doc.add_paragraph("CERTIFICATIONS")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)
            
            for cert in certifications:
                cert_text = f"{cert.get('name', '')} - {cert.get('issuer', '')}"
                if cert.get('date'):
                    cert_text += f" ({cert['date']})"
                doc.add_paragraph(cert_text, style='List Bullet')
            
            doc.add_paragraph()
    
    def _generate_markdown(self, resume_data: Dict[str, Any]) -> str:
        """Generate markdown for PDF conversion"""
        lines = []
        
        # Personal info
        personal = resume_data.get("personal_info", {})
        if personal.get("name"):
            lines.append(f"# {personal['name']}")
            lines.append("")
        
        # Contact
        contact_info = []
        for field in ['email', 'phone', 'location']:
            if personal.get(field):
                contact_info.append(personal[field])
        
        if contact_info:
            lines.append(f"**{' | '.join(contact_info)}**")
            lines.append("")
        
        # Summary
        if resume_data.get("professional_summary"):
            lines.append("## Professional Summary")
            lines.append("")
            lines.append(resume_data["professional_summary"])
            lines.append("")
        
        # Skills
        skills = resume_data.get("core_competencies", [])
        if skills:
            lines.append("## Core Competencies")
            lines.append("")
            lines.append(" • ".join(skills))
            lines.append("")
        
        # Experience
        experience = resume_data.get("professional_experience", [])
        if experience:
            lines.append("## Professional Experience")
            lines.append("")
            
            for job in experience:
                lines.append(f"**{job.get('position', '')} | {job.get('company', '')}**")
                lines.append(f"*{job.get('duration', '')} | {job.get('location', '')}*")
                lines.append("")
                
                for achievement in job.get('achievements', []):
                    lines.append(f"- {achievement}")
                lines.append("")
        
        # Education
        education = resume_data.get("education", [])
        if education:
            lines.append("## Education")
            lines.append("")
            
            for edu in education:
                lines.append(f"**{edu.get('degree', '')} in {edu.get('field', '')}**")
                lines.append(f"*{edu.get('institution', '')} | {edu.get('graduation_year', '')}*")
                lines.append("")
        
        return "\n".join(lines)

def main():
    """Main command-line interface"""
    parser = argparse.ArgumentParser(
        description="Generate tailored resumes from JSON files using Groq LLM",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python resume_generator.py resume.json job.json --format pdf
  python resume_generator.py data/resume.json data/job.json --output tailored_resume --format docx
  python resume_generator.py resume.json job.json --format pdf --api-key gsk_your_key_here
"""
    )
    
    parser.add_argument("resume_json", help="Path to resume data JSON file")
    parser.add_argument("job_json", help="Path to job requirements JSON file")
    parser.add_argument("--output", "-o", default=None,
                       help="Output filename (without extension)")
    parser.add_argument("--format", "-f", choices=["pdf", "docx", "both"],
                       default="pdf", help="Output format (default: pdf)")
    parser.add_argument("--api-key", "-k", default=None,
                       help="Groq API key (or edit code to hardcode)")
    parser.add_argument("--verbose", "-v", action="store_true",
                       help="Enable verbose logging")
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    print("🚀 RESUME GENERATOR - Module 3")
    print("=" * 40)
    
    try:
        # Initialize generator
        generator = CommandLineResumeGenerator(api_key=args.api_key)
        
        # Load JSON files
        print(f"📚 Loading resume data: {args.resume_json}")
        resume_data = generator.load_json_file(args.resume_json)
        
        print(f"📚 Loading job data: {args.job_json}")
        job_data = generator.load_json_file(args.job_json)
        
        # Generate output filename if not provided
        if args.output:
            base_output = args.output
        else:
            candidate_name = resume_data.get("personal_info", {}).get("name", "resume")
            safe_name = "".join(c for c in candidate_name if c.isalnum() or c in (' ', '-', '_')).strip()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_output = f"{safe_name}_tailored_{timestamp}"
        
        # Tailor resume
        print(f"🤖 Tailoring resume using Groq {generator.model}...")
        tailoring_result = generator.tailor_resume(resume_data, job_data)
        
        if not tailoring_result["success"]:
            print(f"❌ TAILORING FAILED: {tailoring_result['error']}")
            sys.exit(1)
        
        tailored_resume = tailoring_result["tailored_resume"]
        
        # Generate documents
        success_count = 0
        
        if args.format in ["docx", "both"]:
            output_path = f"{base_output}.docx"
            print(f"📄 Generating DOCX: {output_path}")
            docx_result = generator.generate_docx(tailored_resume, output_path)
            
            if docx_result["success"]:
                success_count += 1
            else:
                print(f"❌ DOCX generation failed: {docx_result['error']}")
        
        if args.format in ["pdf", "both"]:
            output_path = f"{base_output}.pdf"
            print(f"📄 Generating PDF: {output_path}")
            pdf_result = generator.generate_pdf(tailored_resume, output_path)
            
            if pdf_result["success"]:
                success_count += 1
            else:
                print(f"❌ PDF generation failed: {pdf_result['error']}")
        
        # Summary
        if success_count > 0:
            print(f"\n🎉 SUCCESS! Generated {success_count} file(s)")
            
            # Show optimization info
            candidate_name = tailored_resume.get("personal_info", {}).get("name", "Candidate")
            job_title = job_data.get("job_title", "Position")
            skills_count = len(tailored_resume.get("core_competencies", []))
            
            print(f"\n📊 TAILORING SUMMARY:")
            print(f"   Candidate: {candidate_name}")
            print(f"   Target Job: {job_title}")
            print(f"   Skills Optimized: {skills_count}")
            print(f"   Model Used: {tailoring_result['model']}")
        else:
            print(f"\n❌ FAILED: No files were generated successfully")
            sys.exit(1)
            
    except FileNotFoundError as e:
        print(f"❌ FILE ERROR: {str(e)}")
        sys.exit(1)
    except ValueError as e:
        print(f"❌ JSON ERROR: {str(e)}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ UNEXPECTED ERROR: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()